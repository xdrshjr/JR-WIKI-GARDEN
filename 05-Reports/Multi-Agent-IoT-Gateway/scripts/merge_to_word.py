#!/usr/bin/env python3
"""
多智能体物联网中枢系统技术报告 - 汇总脚本
将多个Markdown文件合并并转换为Word文档
"""

import os
import sys
import re
from datetime import datetime
from pathlib import Path

# 尝试导入python-docx，如果不存在则提示安装
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("请先安装python-docx: pip install python-docx")
    sys.exit(1)

try:
    import markdown
except ImportError:
    print("请先安装markdown: pip install markdown")
    sys.exit(1)

# 配置
REPORTS_DIR = Path(__file__).parent.parent / "reports"
IMAGES_DIR = Path(__file__).parent.parent / "images"
OUTPUT_DIR = Path(__file__).parent.parent

def read_markdown_file(filepath):
    """读取Markdown文件内容"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def convert_md_to_docx():
    """将MD文件转换为Word文档"""
    
    # 定义文件顺序
    md_files = [
        "01-系统架构设计.md",
        "02-通信协议与边端集成.md",
        "03-安全体系设计.md",
        "04-多智能体协作机制.md",
        "05-参考文献.md"
    ]
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    
    # 添加封面
    add_cover_page(doc)
    
    # 添加目录占位符
    doc.add_page_break()
    add_toc_page(doc)
    doc.add_page_break()
    
    # 处理每个MD文件
    for md_file in md_files:
        md_path = REPORTS_DIR / md_file
        if not md_path.exists():
            print(f"警告: 文件不存在 {md_path}")
            continue
        
        print(f"处理: {md_file}")
        content = read_markdown_file(md_path)
        
        # 解析并添加内容
        parse_and_add_content(doc, content, md_file)
        
        # 文件之间添加分页
        doc.add_page_break()
    
    # 保存文档
    output_path = OUTPUT_DIR / "多智能体物联网中枢系统技术报告.docx"
    doc.save(output_path)
    print(f"\n✅ 报告已生成: {output_path}")
    
    return output_path

def add_cover_page(doc):
    """添加封面页"""
    # 添加空行调整位置
    for _ in range(6):
        doc.add_paragraph()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("多智能体物联网中枢系统")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("技术报告")
    run.font.size = Pt(22)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(51, 51, 51)
    
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 描述
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("基于OpenClaw架构的分布式智能物联网平台设计与实现")
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    # 空行
    for _ in range(6):
        doc.add_paragraph()
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"报告日期: {datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'

def add_toc_page(doc):
    """添加目录页"""
    # 目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    
    doc.add_paragraph()
    
    # 手动添加目录项
    toc_items = [
        ("一、系统架构设计", 1),
        ("二、通信协议与边端集成", 1),
        ("三、安全体系设计", 1),
        ("四、多智能体协作机制", 1),
        ("五、参考文献", 1),
    ]
    
    for item, level in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5 * (level - 1))
        run = para.add_run(item)
        run.font.size = Pt(12)
        run.font.name = 'Microsoft YaHei'

def parse_and_add_content(doc, content, filename):
    """解析Markdown内容并添加到Word文档"""
    
    # 简单的Markdown解析
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 处理标题
        if line.startswith('# '):
            # H1 - 章节标题
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(0, 51, 102)
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            i += 1
            continue
        
        elif line.startswith('## '):
            # H2
            text = line[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(0, 76, 153)
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        
        elif line.startswith('### '):
            # H3
            text = line[4:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(51, 51, 51)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # 处理代码块
        elif line.startswith('```'):
            # 开始代码块
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束标记
            
            # 添加代码块
            if code_lines:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(51, 51, 51)
                para.paragraph_format.left_indent = Inches(0.3)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
            continue
        
        # 处理引用块
        elif line.startswith('> '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # 处理表格（简化处理）
        elif '|' in line and ('---' in line or line.strip().startswith('|')):
            # 跳过表格分隔行
            if '---' in line:
                i += 1
                continue
            
            # 收集表格行
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                if '---' not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            
            # 创建表格
            if len(table_lines) >= 1:
                create_table_from_md(doc, table_lines)
            continue
        
        # 处理图片引用
        elif '![' in line and '](' in line:
            # 提取图片路径
            match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text, img_path = match.groups()
                # 尝试找到图片
                img_filename = os.path.basename(img_path)
                img_full_path = IMAGES_DIR / img_filename
                
                if img_full_path.exists():
                    try:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(str(img_full_path), width=Inches(5.5))
                        
                        # 添加图片说明
                        if alt_text:
                            caption = doc.add_paragraph()
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = caption.add_run(alt_text)
                            cap_run.font.size = Pt(10)
                            cap_run.font.italic = True
                            cap_run.font.color.rgb = RGBColor(102, 102, 102)
                    except Exception as e:
                        print(f"  警告: 无法插入图片 {img_filename}: {e}")
                else:
                    # 尝试其他路径
                    alt_paths = [
                        REPORTS_DIR / "images" / img_filename,
                        REPORTS_DIR / img_filename,
                    ]
                    for alt_path in alt_paths:
                        if alt_path.exists():
                            try:
                                para = doc.add_paragraph()
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = para.add_run()
                                run.add_picture(str(alt_path), width=Inches(5.5))
                                
                                if alt_text:
                                    caption = doc.add_paragraph()
                                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cap_run = caption.add_run(alt_text)
                                    cap_run.font.size = Pt(10)
                                    cap_run.font.italic = True
                                    cap_run.font.color.rgb = RGBColor(102, 102, 102)
                            except Exception as e:
                                print(f"  警告: 无法插入图片 {img_filename}: {e}")
                            break
                    else:
                        print(f"  警告: 找不到图片 {img_filename}")
            
            i += 1
            continue
        
        # 处理普通段落
        elif line.strip():
            # 处理行内格式
            para = doc.add_paragraph()
            add_formatted_text(para, line.strip())
            i += 1
            continue
        
        # 空行
        else:
            i += 1

def create_table_from_md(doc, table_lines):
    """从Markdown表格行创建Word表格"""
    if not table_lines:
        return
    
    # 解析第一行获取列数
    first_row = table_lines[0].split('|')
    first_row = [cell.strip() for cell in first_row if cell.strip()]
    num_cols = len(first_row)
    num_rows = len(table_lines)
    
    if num_cols == 0 or num_rows == 0:
        return
    
    # 创建表格
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    # 填充数据
    for row_idx, line in enumerate(table_lines):
        cells = line.split('|')
        cells = [cell.strip() for cell in cells if cell.strip()]
        
        for col_idx, cell_text in enumerate(cells[:num_cols]):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # 表格后添加空行

def add_formatted_text(para, text):
    """添加带格式的文本"""
    # 处理粗体 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体
            run = para.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # 斜体
            run = para.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Microsoft YaHei'
        elif part.startswith('`') and part.endswith('`'):
            # 代码
            run = para.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(204, 0, 0)
        else:
            # 普通文本
            run = para.add_run(part)
            run.font.name = 'Microsoft YaHei'

if __name__ == '__main__':
    print("=" * 60)
    print("多智能体物联网中枢系统技术报告 - 汇总脚本")
    print("=" * 60)
    print()
    
    # 检查依赖
    try:
        output = convert_md_to_docx()
        print("\n" + "=" * 60)
        print("✅ 报告生成成功!")
        print(f"📄 输出文件: {output}")
        print("=" * 60)
    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
