#!/usr/bin/env python3
"""
多智能体物联网中枢系统技术报告 - 改进版汇总脚本
更好的排版和图片处理
"""

import os
import sys
import re
from datetime import datetime
from pathlib import Path

# 导入依赖
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装python-docx")
    sys.exit(1)

# 配置
REPORTS_DIR = Path(__file__).parent.parent / "reports"
IMAGES_DIR = Path(__file__).parent.parent / "images"
OUTPUT_DIR = Path(__file__).parent.parent

def set_chinese_font(run, font_name='Microsoft YaHei', font_size=10.5, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = bold
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_formatted_text(paragraph, text):
    """添加带格式的文本，处理粗体、斜体、行内代码"""
    if not text:
        return
    
    # 处理粗体 **text** 和 __text__
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 粗体
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, bold=True)
        elif part.startswith('__') and part.endswith('__'):
            # 粗体（下划线形式）
            run = paragraph.add_run(part[2:-2])
            set_chinese_font(run, bold=True)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # 斜体
            run = paragraph.add_run(part[1:-1])
            set_chinese_font(run)
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # 行内代码
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(204, 0, 0)
        else:
            # 普通文本
            run = paragraph.add_run(part)
            set_chinese_font(run)

def parse_simple_table(lines):
    """解析简单的Markdown表格"""
    if not lines:
        return None
    
    rows = []
    for line in lines:
        if '|' in line and '---' not in line:
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # 移除空单元格
            if cells:
                rows.append(cells)
    
    if not rows:
        return None
    
    # 确定列数
    num_cols = max(len(row) for row in rows)
    
    # 统一每行的列数
    for row in rows:
        while len(row) < num_cols:
            row.append('')
    
    return rows

def add_cover_page(doc):
    """添加封面页"""
    # 添加空行调整位置
    for _ in range(8):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("多智能体物联网中枢系统")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(0, 51, 102)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("技术报告")
    run.font.size = Pt(26)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(51, 51, 51)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 空行
    for _ in range(4):
        doc.add_paragraph()
    
    # 描述
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("基于OpenClaw架构的分布式智能物联网平台\n设计与实现")
    run.font.size = Pt(14)
    run.font.name = 'Microsoft YaHei'
    run.font.color.rgb = RGBColor(102, 102, 102)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"报告日期：{datetime.now().strftime('%Y年%m月%d日')}")
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_toc_page(doc):
    """添加目录页"""
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目  录")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    doc.add_paragraph()
    
    # 目录项
    toc_items = [
        ("一、系统架构设计", 1),
        ("", 0),
        ("二、通信协议与边端集成", 1),
        ("", 0),
        ("三、安全体系设计", 1),
        ("", 0),
        ("四、多智能体协作机制", 1),
        ("", 0),
        ("五、参考文献", 1),
    ]
    
    for item, level in toc_items:
        if item:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(item)
            set_chinese_font(run, font_size=12)

def find_image_file(img_name):
    """查找图片文件"""
    possible_paths = [
        IMAGES_DIR / img_name,
        REPORTS_DIR / "images" / img_name,
        REPORTS_DIR / img_name,
    ]
    
    for path in possible_paths:
        if path.exists():
            return path
    return None

def convert_md_to_docx():
    """将MD文件转换为Word文档"""
    
    md_files = [
        ("01-系统架构设计.md", "一、系统架构设计"),
        ("02-通信协议与边端集成.md", "二、通信协议与边端集成"),
        ("03-安全体系设计.md", "三、安全体系设计"),
        ("04-多智能体协作机制.md", "四、多智能体协作机制"),
        ("05-参考文献.md", "五、参考文献"),
    ]
    
    # 创建文档
    doc = Document()
    
    # 设置默认段落样式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 添加封面
    add_cover_page(doc)
    doc.add_page_break()
    
    # 添加目录
    add_toc_page(doc)
    doc.add_page_break()
    
    # 处理每个MD文件
    for md_file, chapter_title in md_files:
        md_path = REPORTS_DIR / md_file
        if not md_path.exists():
            print(f"⚠️ 文件不存在: {md_path}")
            continue
        
        print(f"📄 处理: {md_file}")
        
        # 添加章节标题
        chapter_para = doc.add_paragraph()
        chapter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = chapter_para.add_run(chapter_title)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = RGBColor(0, 51, 102)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        chapter_para.paragraph_format.space_before = Pt(24)
        chapter_para.paragraph_format.space_after = Pt(24)
        
        # 读取并解析内容
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        parse_content(doc, content)
        
        # 章节之间分页
        doc.add_page_break()
    
    # 保存
    output_path = OUTPUT_DIR / "多智能体物联网中枢系统技术报告.docx"
    doc.save(output_path)
    print(f"\n✅ 报告已生成: {output_path}")
    return output_path

def parse_content(doc, content):
    """解析Markdown内容"""
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # 结束代码块
                in_code_block = False
                if code_lines:
                    # 添加代码块
                    code_text = '\n'.join(code_lines)
                    para = doc.add_paragraph()
                    run = para.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(51, 51, 51)
                    para.paragraph_format.left_indent = Cm(0.5)
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 处理表格
        if '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # 结束表格
            in_table = False
            rows = parse_simple_table(table_lines)
            if rows and len(rows) > 0:
                num_cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_text
                        # 设置单元格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_chinese_font(run, font_size=9)
                
                doc.add_paragraph()
            table_lines = []
            continue
        
        # 处理标题
        if line.startswith('# '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(0, 76, 153)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            para.paragraph_format.space_before = Pt(20)
            para.paragraph_format.space_after = Pt(10)
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(0, 102, 204)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
            run.font.color.rgb = RGBColor(51, 51, 51)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # 处理图片
        if '![' in line and '](' in line:
            match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                img_name = os.path.basename(img_path)
                
                # 查找图片
                img_file = find_image_file(img_name)
                
                if img_file:
                    try:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        
                        # 插入图片，限制宽度
                        if img_name.endswith('.svg'):
                            print(f"  ⚠️ SVG图片无法直接插入: {img_name}")
                        else:
                            run.add_picture(str(img_file), width=Inches(5.5))
                            
                            # 添加图片说明
                            if alt_text and not alt_text.startswith('../'):
                                caption = doc.add_paragraph()
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cap_run = caption.add_run(alt_text)
                                set_chinese_font(cap_run, font_size=9)
                                cap_run.font.italic = True
                                cap_run.font.color.rgb = RGBColor(102, 102, 102)
                    except Exception as e:
                        print(f"  ⚠️ 无法插入图片 {img_name}: {e}")
                else:
                    print(f"  ⚠️ 找不到图片: {img_name}")
            
            i += 1
            continue
        
        # 处理引用块
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_chinese_font(run)
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
            para.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        
        # 处理普通段落
        if line.strip():
            para = doc.add_paragraph()
            add_formatted_text(para, line.strip())
            para.paragraph_format.space_after = Pt(6)
        
        i += 1

if __name__ == '__main__':
    print("=" * 60)
    print("多智能体物联网中枢系统技术报告 - 改进版汇总")
    print("=" * 60)
    print()
    
    try:
        output = convert_md_to_docx()
        print("\n" + "=" * 60)
        print("✅ 报告生成成功!")
        print(f"📄 输出文件: {output}")
        print("=" * 60)
    except Exception as e:
        print(f"\n❌ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
